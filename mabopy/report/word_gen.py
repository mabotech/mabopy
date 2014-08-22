# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Inches

document = Document()

document.add_heading(u'FT汽车', level=0)

p = document.add_paragraph(u'汽车工程研究院 ')
p.add_run(u'试验中心').bold = True
p.add_run(u'试验数据管理系统')
p.add_run(u'项目二期。').italic = True

document.add_heading(u'报告说明Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='IntenseQuote')

document.add_paragraph(
    u'数据项1 first item in unordered list', style='ListBullet'
)
document.add_paragraph(
    u'有序数据项1 first item in ordered list', style='ListNumber'
)

document.add_picture('foton.png', width=Inches(1.25))

document.add_picture('report1.png', width=Inches(3.25))

document.add_heading(
    u'数据项', level=2
)

table = document.add_table(rows=1, cols=3, style='TableGrid')
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'

for item in xrange(0,4):
    row_cells = table.add_row().cells
    row_cells[0].text = str(item)
    row_cells[1].text = str(item)
    row_cells[2].text = "item.desc"

document.add_page_break()

document.save('foton_test5.docx')